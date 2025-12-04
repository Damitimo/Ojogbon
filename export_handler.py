"""
Export Handler - Exports resumes to various formats (PDF, DOCX, TXT)
"""
from pathlib import Path
from typing import Dict
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


class ResumeExporter:
    """Handles exporting resumes to different formats"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def _add_bottom_border(self, paragraph):
        """Add a bottom border to a paragraph using oxml"""
        try:
            p = paragraph._p  # get the oxml element
            pPr = p.get_or_add_pPr()
            
            # Create bottom border element
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')  # border width (1/8 pt)
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), '000000')  # black
            
            # Create borders element and add bottom border
            borders = OxmlElement('w:pBdr')
            borders.append(bottom_border)
            
            pPr.append(borders)
        except Exception as e:
            # If oxml approach fails, silently continue
            pass
    
    def export_to_pdf(self, resume_data: Dict, filename: str = None) -> Path:
        """Export resume to PDF format"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.pdf"
        
        output_path = self.output_dir / filename
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Set Times New Roman font for the entire document
        pdf.add_font('Times', '', 'times.ttf', uni=True)
        pdf.add_font('Times', 'B', 'timesbd.ttf', uni=True)
        pdf.add_font('Times', 'I', 'timesi.ttf', uni=True)
        pdf.add_font('Times', 'BI', 'timesbi.ttf', uni=True)
        
        # Personal Info
        personal = resume_data.get('personal_info', {})
        pdf.set_font('Times', 'B', 20)
        pdf.cell(0, 10, personal.get('name', 'Your Name'), ln=True, align='C')
        
        pdf.set_font('Times', '', 10)
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])
        
        pdf.cell(0, 5, ' | '.join(contact_info), ln=True, align='C')
        
        links = []
        if personal.get('linkedin'):
            links.append(f"LinkedIn: {personal['linkedin']}")
        if personal.get('github'):
            links.append(f"GitHub: {personal['github']}")
        
        if links:
            pdf.cell(0, 5, ' | '.join(links), ln=True, align='C')
        
        pdf.ln(5)
        
        # Summary
        if resume_data.get('summary'):
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'PROFESSIONAL SUMMARY', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            pdf.set_font('Times', '', 10)
            pdf.multi_cell(0, 5, resume_data['summary'])
            pdf.ln(3)
        
        # Experience
        if resume_data.get('experience'):
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'PROFESSIONAL EXPERIENCE', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            for exp in resume_data['experience']:
                pdf.set_font('Times', 'B', 11)
                pdf.cell(0, 6, f"{exp['title']} - {exp['company']}", ln=True)
                
                pdf.set_font('Times', 'I', 9)
                date_location = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                if exp.get('location'):
                    date_location += f" | {exp['location']}"
                pdf.cell(0, 5, date_location, ln=True)
                
                pdf.set_font('Times', '', 10)
                for bullet in exp.get('description', []):
                    pdf.multi_cell(0, 5, f"  • {bullet}")
                
                pdf.ln(2)
        
        # Projects
        if resume_data.get('projects') and len(resume_data['projects']) > 0:
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'PROJECTS', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            for project in resume_data['projects']:
                pdf.set_font('Times', 'B', 11)
                pdf.cell(0, 6, project['name'], ln=True)
                
                pdf.set_font('Times', 'I', 9)
                if project.get('technologies'):
                    pdf.cell(0, 5, f"Technologies: {', '.join(project['technatives'])}", ln=True)
                
                pdf.set_font('Times', '', 10)
                pdf.multi_cell(0, 5, project['description'])
                
                if project.get('achievements'):
                    for achievement in project['achievements']:
                        pdf.multi_cell(0, 5, f"  • {achievement}")
                
                pdf.ln(2)
        
        # Education
        if resume_data.get('education'):
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'EDUCATION', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            for edu in resume_data['education']:
                pdf.set_font('Times', 'B', 11)
                pdf.cell(0, 6, edu['degree'], ln=True)
                
                pdf.set_font('Times', '', 10)
                edu_info = edu['institution']
                if edu.get('location'):
                    edu_info += f" | {edu['location']}"
                if edu.get('graduation_date'):
                    edu_info += f" | {edu['graduation_date']}"
                pdf.cell(0, 5, edu_info, ln=True)
                
                if edu.get('gpa'):
                    pdf.cell(0, 5, f"GPA: {edu['gpa']}", ln=True)
                
                # Add achievements as bullets if they exist
                if edu.get('achievements') and len(edu['achievements']) > 0:
                    for achievement in edu['achievements']:
                        pdf.multi_cell(0, 5, f"  • {achievement.lstrip('• ')}")
                
                pdf.ln(2)
        
        # Skills
        if resume_data.get('skills'):
            pdf.ln(5)
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'SKILLS', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            pdf.set_font('Times', 'B', 10)
            skills = resume_data['skills']
            
            if skills.get('technical'):
                pdf.cell(40, 5, 'Technical:', 0, 0)
                pdf.set_font('Times', '', 10)
                pdf.multi_cell(0, 5, ', '.join(skills['technical']))
                pdf.ln(1)
                
            if skills.get('languages'):
                pdf.set_font('Times', 'B', 10)
                pdf.cell(40, 5, 'Languages:', 0, 0)
                pdf.set_font('Times', '', 10)
                pdf.multi_cell(0, 5, ', '.join(skills['languages']))
                pdf.ln(1)
                
            if skills.get('tools'):
                pdf.set_font('Times', 'B', 10)
                pdf.cell(40, 5, 'Tools:', 0, 0)
                pdf.set_font('Times', '', 10)
                pdf.multi_cell(0, 5, ', '.join(skills['tools']))
        
        # Certifications
        if resume_data.get('certifications') and len(resume_data['certifications']) > 0:
            pdf.ln(2)
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 7, 'CERTIFICATIONS', ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            pdf.set_font('Times', '', 10)
            for cert in resume_data['certifications']:
                pdf.cell(0, 5, f"• {cert}", ln=True)
        
        pdf.output(str(output_path))
        return output_path
    
    def export_to_docx(self, resume_data: Dict, filename: str = None) -> Path:
        """Export resume to DOCX format"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.docx"
        
        output_path = self.output_dir / filename
        
        doc = Document()
        
        # Set page margins (0.5 inches on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Set default font to Times New Roman, 9.5pt
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(9.5)
        
        # Set No Spacing style to Times New Roman
        try:
            no_spacing_style = doc.styles['No Spacing']
            no_spacing_style.font.name = 'Times New Roman'
            no_spacing_style.font.size = Pt(9.5)
        except:
            pass
        
        # Set heading styles to Times New Roman, 9.5pt
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(9.5)
            style.font.color.rgb = None  # Remove any color
        
        # Personal Info
        personal = resume_data.get('personal_info', {})
        
        # Create a table with one cell for full-width name with bottom border
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False
        table.width = Inches(6.5)  # Full width minus margins
        
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add name (remove underline, increase to 16pt, all caps)
        name_run = cell.paragraphs[0].add_run(personal.get('name', 'Your Name').upper())
        name_run.bold = True
        name_run.font.size = Pt(16)  # Increased to 16pt
        name_run.font.name = 'Times New Roman'
        # Removed underline
        
        # Add bottom border to name
        self._add_bottom_border(cell.paragraphs[0])
        
        # Reduce space after name paragraph to match other spacing
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        
        # Combine all contact info into a single line with reduced spacing
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin'):
            linkedin_url = personal['linkedin'].replace('http://', '').replace('https://', '').replace('www.', '')
            contact_parts.append(linkedin_url)
        if personal.get('github'):
            github_url = personal['github'].replace('http://', '').replace('https://', '').replace('www.', '')
            contact_parts.append(github_url)
        
        if contact_parts:
            contact = doc.add_paragraph(' | '.join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_after = Pt(2)  # Reduced to match other spacing
        
        # Summary with reduced spacing
        if resume_data.get('summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY', style='Heading 2')
            summary_heading.runs[0].font.name = 'Times New Roman'
            summary_heading.runs[0].font.size = Pt(9.5)
            summary_heading.runs[0].bold = True
            
            # Add bottom border
            self._add_bottom_border(summary_heading)
            
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.paragraph_format.space_after = Pt(2)  # Consistent spacing
            
        # Education
        if resume_data.get('education'):
            edu_heading = doc.add_paragraph('EDUCATION', style='Heading 2')
            edu_heading.runs[0].font.name = 'Times New Roman'
            edu_heading.runs[0].font.size = Pt(9.5)
            edu_heading.runs[0].bold = True
            
            # Add bottom border
            self._add_bottom_border(edu_heading)
            
            for edu in resume_data['education']:
                # First line: School name (bold) - Location (regular) with date right-aligned
                p = doc.add_paragraph(style='No Spacing')
                p.paragraph_format.space_before = Pt(5)  # Add 5pt space before each education
                p.paragraph_format.space_after = Pt(0)
                
                # Institution name in bold
                institution = p.add_run(edu.get('institution', ''))
                institution.bold = True
                institution.font.name = 'Times New Roman'
                
                # Add location with hyphen in regular font (ensure Times New Roman)
                if edu.get('location'):
                    location_run = p.add_run(f" - {edu['location']}")
                    location_run.font.name = 'Times New Roman'
                
                # Add right-aligned graduation date (bold to match experience)
                if edu.get('end_date'):
                    date_run = p.add_run("\t" + edu['end_date'])
                    date_run.bold = True
                    date_run.font.name = 'Times New Roman'
                    tab_stops = p.paragraph_format.tab_stops
                    # Match the same tab stops as experience section
                    tab_stop = tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
                    tab_stop = tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
                
                # Second line: Degree in italics
                degree_para = doc.add_paragraph(style='No Spacing')
                degree_para.paragraph_format.space_after = Pt(0)  # Remove extra space after degree
                
                degree = degree_para.add_run(edu.get('degree', ''))
                degree.italic = True
                degree.font.name = 'Times New Roman'
                
                # Add achievements as bullets if they exist
                if edu.get('achievements') and len(edu['achievements']) > 0:
                    for achievement in edu['achievements']:
                        achievement_para = doc.add_paragraph(style='List Bullet')
                        achievement_para.paragraph_format.left_indent = Inches(0.5)
                        achievement_para.paragraph_format.space_before = Pt(0)
                        achievement_para.paragraph_format.space_after = Pt(0)
                        achievement_run = achievement_para.add_run(achievement.lstrip('• '))
                        achievement_run.font.name = 'Times New Roman'

        # Experience
        if resume_data.get('experience'):
            exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 2')
            exp_heading.runs[0].font.name = 'Times New Roman'
            exp_heading.runs[0].font.size = Pt(9.5)
            exp_heading.runs[0].bold = True
            
            # Add bottom border
            self._add_bottom_border(exp_heading)

            for exp in resume_data['experience']:
                # Company, location, and date on one line
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)  # Add 5pt space before each experience
                p.paragraph_format.space_after = Pt(0)
                
                # Company name in bold
                company = p.add_run(f"{exp['company']}")
                company.bold = True
                
                # Add location if exists (ensure Times New Roman)
                if exp.get('location'):
                    location_run = p.add_run(f" - {exp['location']}")
                    location_run.font.name = 'Times New Roman'
                
                # Add right-aligned date with tab
                p.add_run("\t")
                date_str = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                date = p.add_run(date_str)
                date.bold = True
                
                # Set tab stop for right alignment of the date
                tab_stops = p.paragraph_format.tab_stops
                # Add tab stop at the right margin (8.5" - 1" margins = 7.5")
                tab_stop = tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
                # Add a left tab stop to keep the left part left-aligned
                tab_stop = tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
                
                # Job title in italic on next line with Times New Roman
                title_para = doc.add_paragraph(style='No Spacing')
                title_para.paragraph_format.space_after = Pt(0)
                title = title_para.add_run(exp.get('title', ''))
                title.italic = True
                title.font.name = 'Times New Roman'
                
                # Description bullets
                for bullet in exp.get('description', []):
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
                    bullet_para.paragraph_format.space_before = Pt(0)
                    bullet_para.paragraph_format.space_after = Pt(0)
                    bullet_para.add_run(bullet)
        # Projects
        if resume_data.get('projects') and len(resume_data['projects']) > 0:
            proj_heading = doc.add_paragraph('PROJECTS', style='Heading 2')
            proj_heading.runs[0].font.name = 'Times New Roman'
            proj_heading.runs[0].font.size = Pt(9.5)
            proj_heading.runs[0].bold = True
            
            # Add bottom border
            self._add_bottom_border(proj_heading)
            
            for project in resume_data['projects']:
                # Project name in bold
                p = doc.add_paragraph(style='No Spacing')
                p.paragraph_format.space_after = Pt(0)
                name = p.add_run(project.get('name', ''))
                name.bold = True
                name.font.name = 'Times New Roman'
                
                # Technologies in italic
                if project.get('technologies'):
                    tech_para = doc.add_paragraph(style='No Spacing')
                    tech_para.paragraph_format.space_after = Pt(0)
                    tech = tech_para.add_run('Technologies: ')
                    tech.italic = True
                    tech.font.name = 'Times New Roman'
                    tech_text = tech_para.add_run(', '.join(project['technologies']))
                    tech_text.font.name = 'Times New Roman'
                
                # Project description
                if project.get('description'):
                    desc_para = doc.add_paragraph(style='No Spacing')
                    desc_para.paragraph_format.space_after = Pt(0)
                    desc_run = desc_para.add_run(project['description'])
                    desc_run.font.name = 'Times New Roman'
                
                # Project achievements
                if project.get('achievements'):
                    for achievement in project['achievements']:
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        bullet_para.paragraph_format.left_indent = Inches(0.5)
                        bullet_para.paragraph_format.space_before = Pt(0)
                        bullet_para.paragraph_format.space_after = Pt(0)
                        bullet_run = bullet_para.add_run(achievement)
                        bullet_run.font.name = 'Times New Roman'
                
                # Add small space between projects (5pt)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(5)
        
        # Skills
        if resume_data.get('skills'):
            skills_heading = doc.add_paragraph('SKILLS', style='Heading 2')
            skills_heading.runs[0].font.name = 'Times New Roman'
            skills_heading.runs[0].font.size = Pt(9.5)
            skills_heading.runs[0].bold = True
            
            # Add bottom border
            self._add_bottom_border(skills_heading)
            
            skills = resume_data['skills']
            
            if skills.get('technical'):
                skills_para = doc.add_paragraph(style='No Spacing')
                tech_label = skills_para.add_run('Technical: ')
                tech_label.bold = True
                tech_label.font.name = 'Times New Roman'
                tech_text = skills_para.add_run(', '.join(skills['technical']))
                tech_text.font.name = 'Times New Roman'
                skills_para.paragraph_format.space_after = Pt(0)
                
            if skills.get('languages'):
                skills_para = doc.add_paragraph(style='No Spacing')
                lang_label = skills_para.add_run('Languages: ')
                lang_label.bold = True
                lang_label.font.name = 'Times New Roman'
                lang_text = skills_para.add_run(', '.join(skills['languages']))
                lang_text.font.name = 'Times New Roman'
                skills_para.paragraph_format.space_after = Pt(0)
                
            if skills.get('tools'):
                skills_para = doc.add_paragraph(style='No Spacing')
                tools_label = skills_para.add_run('Tools: ')
                tools_label.bold = True
                tools_label.font.name = 'Times New Roman'
                tools_text = skills_para.add_run(', '.join(skills['tools']))
                tools_text.font.name = 'Times New Roman'
                skills_para.paragraph_format.space_after = Pt(0)
        
        # Certifications
        if resume_data.get('certifications') and len(resume_data['certifications']) > 0:
            cert_heading = doc.add_paragraph('CERTIFICATIONS', style='Heading 2')
            cert_heading.runs[0].font.name = 'Times New Roman'
            cert_heading.runs[0].font.size = Pt(9.5)
            cert_heading.runs[0].bold = True
            
            for cert in resume_data['certifications']:
                cert_para = doc.add_paragraph(style='List Bullet')
                cert_para.paragraph_format.left_indent = Inches(0.5)
                cert_para.paragraph_format.space_before = Pt(0)
                cert_para.paragraph_format.space_after = Pt(0)
                cert_run = cert_para.add_run(cert)
                cert_run.font.name = 'Times New Roman'
        
        # Save the document
        doc.save(str(output_path))
        return output_path
    
    def export_to_txt(self, resume_data: Dict, filename: str = None) -> Path:
        """Export resume to plain text format"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.txt"
        
        output_path = self.output_dir / filename
        
        lines = []
        
        # Personal Info
        personal = resume_data.get('personal_info', {})
        lines.append(personal.get('name', 'Your Name').upper())
        lines.append('=' * 80)
        
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])
        
        if contact_info:
            lines.append(' | '.join(contact_info))
        
        if personal.get('linkedin'):
            lines.append(f"LinkedIn: {personal['linkedin']}")
        if personal.get('github'):
            lines.append(f"GitHub: {personal['github']}")
        
        lines.append('')
        
        # Summary
        if resume_data.get('summary'):
            lines.append('PROFESSIONAL SUMMARY')
            lines.append('-' * 80)
            lines.append(resume_data['summary'])
            lines.append('')
        
        # Experience
        if resume_data.get('experience'):
            lines.append('PROFESSIONAL EXPERIENCE')
            lines.append('-' * 80)
            
            for exp in resume_data['experience']:
                lines.append(f"{exp['title']} - {exp['company']}")
                date_location = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                if exp.get('location'):
                    date_location += f" | {exp['location']}"
                lines.append(date_location)
                
                for bullet in exp.get('description', []):
                    lines.append(f"  • {bullet}")
                
                lines.append('')
        
        # Projects
        if resume_data.get('projects') and len(resume_data['projects']) > 0:
            lines.append('PROJECTS')
            lines.append('-' * 80)
            
            for project in resume_data['projects']:
                lines.append(project['name'])
                if project.get('technologies'):
                    lines.append(f"Technologies: {', '.join(project['technologies'])}")
                lines.append(project['description'])
                
                if project.get('achievements'):
                    for achievement in project['achievements']:
                        lines.append(f"  • {achievement}")
                
                lines.append('')
        
        # Education
        if resume_data.get('education'):
            lines.append('EDUCATION')
            lines.append('-' * 80)
            
            for edu in resume_data['education']:
                lines.append(edu['degree'])
                edu_info = edu['institution']
                if edu.get('location'):
                    edu_info += f" | {edu['location']}"
                if edu.get('graduation_date'):
                    edu_info += f" | {edu['graduation_date']}"
                lines.append(edu_info)
                
                if edu.get('gpa'):
                    lines.append(f"GPA: {edu['gpa']}")
                
                # Add achievements as bullets if they exist
                if edu.get('achievements') and len(edu['achievements']) > 0:
                    for achievement in edu['achievements']:
                        lines.append(f"  • {achievement.lstrip('• ')}")
                
                lines.append('')
        
        # Skills
        if resume_data.get('skills'):
            lines.append('SKILLS')
            lines.append('-' * 80)
            skills = resume_data['skills']
            
            if skills.get('technical'):
                lines.append(f"Technical: {', '.join(skills['technical'])}")
            if skills.get('languages'):
                lines.append(f"Languages: {', '.join(skills['languages'])}")
            if skills.get('tools'):
                lines.append(f"Tools: {', '.join(skills['tools'])}")
            
            lines.append('')
        
        # Certifications
        if resume_data.get('certifications') and len(resume_data['certifications']) > 0:
            lines.append('CERTIFICATIONS')
            lines.append('-' * 80)
            for cert in resume_data['certifications']:
                lines.append(f"• {cert}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return output_path
